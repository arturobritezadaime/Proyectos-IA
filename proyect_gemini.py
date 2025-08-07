from config import cargar_api_key
from modelos import obtener_modelos, mostrar_modelos
from entrada_usuario import seleccionar_modelo
from generador import generar_contenido
from rich.console import Console
import csv
from docx import Document

console = Console()

# ================================
# 🔹 Función principal (modo conversacional)
# ================================

def main():
    try:
        cargar_api_key()
        modelos = obtener_modelos()
        mostrar_modelos(modelos)

        modelo_info = seleccionar_modelo(modelos)
        if not modelo_info:
            raise ValueError("❌ Número de modelo no válido.")

        modelo_id = modelo_info["id"]
        historial_conversacion = []

        console.print("\n💬 [bold cyan]Modo conversacional activado.[/bold cyan] Escribí 'salir' para terminar.\n")

        while True:
            prompt = input("Tú: ").strip()
            if prompt.lower() in ["salir", "exit", "quit"]:
                console.print("👋 [bold yellow]Fin de la conversación.[/bold yellow]")
                break

            historial_conversacion.append({"role": "user", "parts": [prompt]})
            respuesta = generar_contenido(modelo_id, historial_conversacion)
            console.print(f"\n🤖 [bold green]Modelo:[/bold green] {respuesta}\n")
            historial_conversacion.append({"role": "model", "parts": [respuesta]})

    except Exception as e:
        console.print(f"\n⚠️ [bold red]Error:[/bold red] {e}", style="red")

    # Guardar conversación si hay historial
    if historial_conversacion:
        guardar = input("\n💾 ¿Querés guardar la conversación? (sí/no): ").strip().lower()
        if guardar in ["sí", "si", "s", "yes", "y"]:
            formato = input("📂 Elegí formato (txt, docx, csv): ").strip().lower()
            formatos_validos = ["txt", "docx", "csv"]

            if formato not in formatos_validos:
                console.print(f"\n⚠️ [yellow]Formato '{formato}' no reconocido. Se guardará en 'txt' por defecto.[/yellow]")
                formato = "txt"

            nombre_archivo = f"conversacion.{formato}"

            if formato == "txt":
                with open(nombre_archivo, "w", encoding="utf-8") as f:
                    for turno in historial_conversacion:
                        f.write(f"{turno['role'].capitalize()}: {turno['parts'][0]}\n\n")

            elif formato == "docx":
                doc = Document()
                doc.add_heading('Conversación con Gemini', level=1)
                for turno in historial_conversacion:
                    doc.add_paragraph(f"{turno['role'].capitalize()}: {turno['parts'][0]}")
                doc.save(nombre_archivo)

            elif formato == "csv":
                with open(nombre_archivo, "w", newline="", encoding="utf-8") as f:
                    writer = csv.writer(f)
                    writer.writerow(["Rol", "Contenido"])
                    for turno in historial_conversacion:
                        writer.writerow([turno["role"], turno["parts"][0]])

            console.print(f"\n✅ [green]Conversación guardada como '{nombre_archivo}'[/green]")


# ================================
# 🔹 Ejecutar
# ================================

if __name__ == "__main__":
    main()
